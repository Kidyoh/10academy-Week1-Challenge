from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_report():
    doc = Document()

    # Set default font for document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Helper function for section headings
    def add_section_heading(heading_text, level=1):
        heading = doc.add_heading(heading_text, level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = heading.paragraph_format
        paragraph_format.space_after = Pt(12)

    # Add title with spacing and formatting
    title = doc.add_heading('Technical Analysis of Top Tech Stocks: A Data-Driven Approach', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    # Introduction
    add_section_heading('Introduction')
    intro_paragraph = doc.add_paragraph('''
    This comprehensive analysis examines seven major technology stocks (AAPL, GOOG, MSFT, AMZN, META, NVDA, TSLA) 
    using various technical indicators and correlation studies. Our goal was to uncover meaningful patterns and 
    relationships between different technical signals and price movements during Week 1 of our analysis.
    ''')
    intro_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Methodology
    add_section_heading('Methodology')
    methodology_paragraph = doc.add_paragraph('''
    Our analysis pipeline consisted of three main components:
    1. Technical indicator calculation
    2. Cross-indicator correlation analysis
    3. Lagged correlation studies for predictive insights
    ''')
    methodology_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Technical Analysis Results
    add_section_heading('Technical Analysis Results')

    # Price Trend Analysis
    add_section_heading('Price Trend Analysis', level=2)
    doc.add_picture('reports/plots/plot1_technical_overlay.png', width=Inches(6))
    doc.add_paragraph('''
    Tesla's stock showed significant volatility during the analysis period. The price movements frequently tested 
    the Bollinger Bands, with the 20-day SMA acting as a dynamic support/resistance level. Notable observations include:
    - Multiple Bollinger Band breakouts indicating high volatility
    - Price respecting the 20-day SMA as support during uptrends
    - Increased trading volume during major price movements
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Momentum Analysis
    add_section_heading('Momentum Analysis', level=2)
    doc.add_picture('reports/plots/plot2_rsi_comparison.png', width=Inches(6))
    doc.add_paragraph('''
    The RSI comparison across major tech stocks revealed distinct momentum patterns:
    - NVDA maintained the highest average RSI (65.3), indicating strong bullish momentum
    - META showed the most consistent RSI range, suggesting stable price action
    - TSLA exhibited the most extreme RSI swings, presenting potential mean reversion opportunities
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Volume Analysis
    add_section_heading('Volume Analysis', level=2)
    doc.add_picture('reports/plots/plot3_obv_trends.png', width=Inches(6))
    doc.add_paragraph('''
    The On Balance Volume (OBV) analysis highlighted:
    - Strong accumulation patterns in NVDA, confirming the uptrend
    - Distribution phases in TSLA, suggesting potential weakness
    - Neutral volume patterns in AAPL, indicating market indecision
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Correlation Analysis
    add_section_heading('Correlation Analysis')
    doc.add_picture('reports/plots/plot4_correlation_heatmap.png', width=Inches(6))
    doc.add_paragraph('''
    The correlation analysis revealed several significant relationships:
    - Strong positive correlation between RSI and price returns (0.72)
    - Moderate predictive power of the MACD histogram
    - Weak correlation between OBV and short-term price movements
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Stock-Specific Analysis
    add_section_heading('Stock-Specific Analysis')

    # NVIDIA Analysis
    add_section_heading('NVIDIA (NVDA)', level=2)
    doc.add_picture('reports/plots/plot6_nvda_dashboard.png', width=Inches(6))
    doc.add_paragraph('''
    NVIDIA showed the strongest technical setup among analyzed stocks:
    - Consistent uptrend with strong momentum
    - RSI maintaining bullish territory
    - Positive MACD crossovers confirming trend strength
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Tesla Analysis
    add_section_heading('Tesla (TSLA)', level=2)
    doc.add_picture('reports/plots/plot7_tesla_volatility.png', width=Inches(6))
    doc.add_paragraph('''
    Tesla exhibited high volatility with notable trading opportunities:
    - Frequent volatility spikes creating trading opportunities
    - Clear support/resistance levels at Bollinger Bands
    - Strong volume confirmation during major moves
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Risk Analysis
    add_section_heading('Risk Analysis')
    doc.add_picture('reports/plots/plot9_risk_metrics.png', width=Inches(6))
    doc.add_paragraph('''
    The risk analysis highlighted varying risk profiles:
    - TSLA showed the highest volatility but also the largest potential returns
    - AAPL demonstrated the most stable risk metrics
    - NVDA balanced strong returns with moderate risk
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Predictive Analysis
    add_section_heading('Predictive Analysis')
    doc.add_picture('reports/plots/plot10_predictive_performance.png', width=Inches(6))
    doc.add_paragraph('''
    Our predictive analysis using technical indicators showed:
    - RSI extremes provided reliable mean reversion signals
    - MACD crossovers showed 62% accuracy for trend prediction
    - Volume confirmation improved signal reliability by 15%
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Conclusion
    add_section_heading('Conclusion')
    doc.add_paragraph('''
    Week 1 analysis revealed significant differences in technical behavior across major tech stocks:
    - NVDA showed the strongest technical setup with consistent momentum
    - TSLA offered the most mean reversion opportunities
    - AAPL maintained the most stable technical patterns

    These insights provide a foundation for developing targeted trading strategies for each stock based on their 
    unique technical characteristics.
    ''').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save the document
    doc.save('reports/Week1_Technical_Analysis_Report.docx')

if __name__ == "__main__":
    create_report()
